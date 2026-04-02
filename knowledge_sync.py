"""Knowledge base ingestion and synchronization helpers."""
from __future__ import annotations

import hashlib
import io
import json
import re
import tempfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List

from bilibili_api import comment, hot, sync, video, video_zone

from knowledge_base import Document, add_document, delete_documents, document_exists

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None


COMMENT_STOPWORDS = {
    "这个",
    "真的",
    "就是",
    "感觉",
    "视频",
    "评论",
    "他们",
    "你们",
    "我们",
    "一个",
    "不是",
    "没有",
    "还是",
    "因为",
    "看到",
    "什么",
    "怎么",
    "真的好",
    "哈哈",
    "哈哈哈",
    "可以",
    "一下",
}
POPULAR_RANK_PARTITIONS = (
    {"label": "番剧", "tid": 13, "board_url": "https://www.bilibili.com/anime/"},
    {"label": "国创", "tid": 167, "board_url": "https://www.bilibili.com/guochuang/"},
    {"label": "纪录片", "tid": 177, "board_url": "https://www.bilibili.com/documentary/"},
    {"label": "动画", "tid": 1, "board_url": "https://www.bilibili.com/v/douga/"},
    {"label": "游戏", "tid": 4, "board_url": "https://www.bilibili.com/v/game/"},
    {"label": "鬼畜", "tid": 119, "board_url": "https://www.bilibili.com/v/kichiku/"},
    {"label": "音乐", "tid": 3, "board_url": "https://www.bilibili.com/v/music"},
    {"label": "舞蹈", "tid": 129, "board_url": "https://www.bilibili.com/v/dance/"},
    {"label": "知识", "tid": 36, "board_url": "https://www.bilibili.com/v/knowledge/"},
    {"label": "科技", "tid": 188, "board_url": "https://www.bilibili.com/v/tech/"},
    {"label": "汽车", "tid": 223, "board_url": "https://www.bilibili.com/v/car"},
    {"label": "运动", "tid": 234, "board_url": "https://www.bilibili.com/v/sports"},
    {"label": "动物圈", "tid": 217, "board_url": "https://www.bilibili.com/v/animal"},
    {"label": "生活", "tid": 160, "board_url": "https://www.bilibili.com/v/life"},
    {"label": "娱乐", "tid": 5, "board_url": "https://www.bilibili.com/v/ent/"},
    {"label": "影视", "tid": 181, "board_url": "https://www.bilibili.com/v/cinephile"},
    {"label": "电影", "tid": 23, "board_url": "https://www.bilibili.com/movie/"},
    {"label": "电视剧", "tid": 11, "board_url": "https://www.bilibili.com/tv/"},
    {"label": "时尚", "tid": 155, "board_url": "https://www.bilibili.com/v/fashion"},
)
LEGACY_PRIMARY_PARTITIONS = ("knowledge", "tech", "life", "game", "ent")
ProgressCallback = Callable[[Dict[str, Any]], None]


def notify_progress(progress_callback: ProgressCallback | None, payload: Dict[str, Any]) -> None:
    if progress_callback is None:
        return
    try:
        progress_callback(payload)
    except Exception:
        return


def normalize_kb_text(text: str) -> str:
    clean = str(text or "").replace("\r", "\n")
    clean = re.sub(r"\n{3,}", "\n\n", clean)
    clean = re.sub(r"[ \t]{2,}", " ", clean)
    return clean.strip()


def safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value or 0)
    except Exception:
        return default


def parse_duration_seconds(value: Any) -> int:
    if value is None:
        return 0
    if isinstance(value, (int, float)):
        return int(value)

    text = str(value or "").strip()
    if not text:
        return 0
    if text.isdigit():
        return int(text)
    if ":" not in text:
        return safe_int(text, 0)

    parts = [segment.strip() for segment in text.split(":")]
    if not parts or any(not part.isdigit() for part in parts):
        return 0

    seconds = 0
    for part in parts:
        seconds = seconds * 60 + int(part)
    return seconds


def keyword_tokens(text: str) -> List[str]:
    return re.findall(r"[\u4e00-\u9fff]{2,8}|[A-Za-z0-9]{2,20}", str(text or "").lower())


def detect_file_suffix(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def read_uploaded_file_content(filename: str, raw_bytes: bytes) -> str:
    suffix = detect_file_suffix(filename)
    if suffix in {".txt", ".md"}:
        return normalize_kb_text(raw_bytes.decode("utf-8", errors="ignore"))
    if suffix == ".docx":
        if DocxDocument is None:
            raise RuntimeError("未安装 python-docx，无法读取 docx 文件。")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_file.write(raw_bytes)
            temp_path = Path(temp_file.name)
        try:
            document = DocxDocument(str(temp_path))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
            return normalize_kb_text(text)
        finally:
            temp_path.unlink(missing_ok=True)
    if suffix == ".pdf":
        if PdfReader is None:
            raise RuntimeError("未安装 pypdf，无法读取 pdf 文件。")
        reader = PdfReader(io.BytesIO(raw_bytes))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        return normalize_kb_text(text)
    raise ValueError("暂不支持该文件格式，仅支持 txt / md / docx / pdf。")


def ingest_uploaded_file(filename: str, raw_bytes: bytes, metadata: Dict[str, Any] | None = None) -> Dict[str, Any]:
    text = read_uploaded_file_content(filename, raw_bytes)
    if not text:
        raise ValueError("文件内容为空，无法写入知识库。")
    content_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
    document_id = f"file:{content_hash}"
    existed = document_exists(metadata_filter={"source": "uploaded_file", "content_hash": content_hash})
    if existed:
        delete_documents(metadata_filter={"source": "uploaded_file", "content_hash": content_hash})
    result = add_document(
        Document(
            id=document_id,
            text=text,
            metadata={
                "source": "uploaded_file",
                "filename": Path(filename).name,
                "content_hash": content_hash,
                **(metadata or {}),
            },
        )
    )
    result["filename"] = Path(filename).name
    result["content_hash"] = content_hash
    return result


def _safe_sync(coro, default):
    try:
        return sync(coro)
    except Exception:
        return default


def _comment_hotwords(aid: int, limit: int = 8) -> List[str]:
    payload = _safe_sync(
        comment.get_comments(aid, comment.CommentResourceType.VIDEO, page_index=1, order=comment.OrderType.LIKE),
        {},
    )
    replies = payload.get("replies") or []
    words: Counter[str] = Counter()
    for item in replies[:15]:
        message = ((item.get("content") or {}).get("message") or "").strip()
        for token in keyword_tokens(message):
            if token in COMMENT_STOPWORDS or len(token) < 2:
                continue
            words[token] += 1
    return [word for word, _ in words.most_common(limit)]


def _format_pub_slot(pub_ts: int) -> str:
    if pub_ts <= 0:
        return "未知"
    dt = datetime.fromtimestamp(pub_ts)
    if 6 <= dt.hour < 10:
        period = "早间"
    elif 10 <= dt.hour < 14:
        period = "午间"
    elif 14 <= dt.hour < 18:
        period = "下午"
    elif 18 <= dt.hour < 23:
        period = "晚高峰"
    else:
        period = "夜间"
    return f"{dt.strftime('%Y-%m-%d %H:%M')}（{period}）"


def _title_advantage(title: str) -> str:
    text = title or ""
    traits = []
    if any(token in text for token in ["？", "?", "为什么", "怎么", "有何", "到底"]):
        traits.append("标题自带提问或冲突，容易激发点击欲")
    if re.search(r"\d+", text):
        traits.append("标题里有数字或期数，记忆点明确")
    if any(token in text for token in ["反差", "坦白局", "实测", "对比", "第一次"]):
        traits.append("标题带有强反差或真实体验信号，进入动机更强")
    if not traits:
        traits.append("标题信息密度高，能快速说明题材和看点")
    return "；".join(traits[:2])


def _script_advantage(duration: int, desc: str) -> str:
    if duration and duration <= 90:
        return "内容长度偏短，通常适合快速抛钩子后直接进入核心信息。"
    if duration and duration <= 240:
        return "中短视频时长，适合按开场钩子 + 两段核心内容 + 互动结尾组织脚本。"
    if desc:
        return "简介里有明确主题延展空间，适合拆成多层信息递进来讲。"
    return "视频适合先立主题，再用中段补充案例或体验，最后用互动问题收束。"


def _rhythm_advantage(duration: int, view: int, like: int) -> str:
    if duration and duration < 60:
        return "整体节奏偏快，更适合连续高信息密度输出。"
    if like and view and like / max(view, 1) > 0.05:
        return "点赞率较高，说明情绪点或节奏点命中较准。"
    return "节奏大概率靠主题推进和镜头切换维持，不依赖单一爆点。"


def _interaction_advantage(title: str, reply: int, hotwords: List[str]) -> str:
    if reply > 500:
        return "评论量较高，说明话题天然适合引发站队、补充经历或观点讨论。"
    if hotwords:
        return f"评论热词集中在「{'、'.join(hotwords[:4])}」，说明互动点已经比较明确。"
    if any(token in title for token in ["你会", "你们", "有没有", "会不会"]):
        return "标题本身自带代入式提问，评论区更容易被调动起来。"
    return "题材具备讨论空间，适合在结尾主动抛出选择题或经历型提问。"


def _tag_advantage(tags: Iterable[str], partition_label: str) -> str:
    clean_tags = [str(tag).strip() for tag in tags if str(tag).strip()]
    if clean_tags:
        return f"标签围绕「{'、'.join(clean_tags[:5])}」展开，利于平台理解内容语义和同类召回。"
    return f"分区语义集中在「{partition_label}」，建议继续围绕同赛道关键词做标签补强。"


def _structured_video_text(board_type: str, item: dict, detail: dict, tags: List[str], hotwords: List[str], board_url: str = "") -> str:
    stat = detail.get("stat") or item.get("stat") or {}
    title = str(detail.get("title") or item.get("title") or "").strip()
    owner = (detail.get("owner") or item.get("owner") or {})
    partition = str(detail.get("tname") or item.get("tname") or "").strip()
    duration = parse_duration_seconds(detail.get("duration") or item.get("duration") or 0)
    view = safe_int(stat.get("view") or item.get("view") or 0)
    like = safe_int(stat.get("like") or item.get("like") or 0)
    reply = safe_int(stat.get("reply") or item.get("reply") or 0)
    pub_ts = safe_int(detail.get("pubdate") or item.get("pubdate") or 0)
    advantages = {
        "标题亮点": _title_advantage(title),
        "脚本结构": _script_advantage(duration, str(detail.get("desc") or "")),
        "节奏把控": _rhythm_advantage(duration, view, like),
        "互动设计": _interaction_advantage(title, reply, hotwords),
        "标签策略": _tag_advantage(tags, partition or "综合"),
    }
    payload = {
        "榜单来源": board_type,
        "视频标题": title,
        "UP主": str(owner.get("name") or item.get("author") or "").strip(),
        "分区": partition,
        "播放量": view,
        "点赞量": like,
        "评论热词": hotwords,
        "核心优点": advantages,
        "发布时间点": _format_pub_slot(pub_ts),
        "BVID": str(detail.get("bvid") or item.get("bvid") or "").strip(),
        "榜单链接": str(board_url or "").strip(),
        "链接": f"https://www.bilibili.com/video/{detail.get('bvid') or item.get('bvid')}",
    }
    return normalize_kb_text(json.dumps(payload, ensure_ascii=False, indent=2))


def _video_detail(bvid: str) -> dict:
    if not bvid:
        return {}
    return _safe_sync(video.Video(bvid=bvid).get_info(), {})


def _video_tags(video_obj: dict) -> List[str]:
    bvid = str(video_obj.get("bvid") or "").strip()
    if not bvid:
        return []
    data = _safe_sync(video.Video(bvid=bvid).get_tags(), [])
    if not isinstance(data, list):
        return []
    return [str(item.get("tag_name") or "").strip() for item in data if str(item.get("tag_name") or "").strip()]


def _ingest_hot_items(
    board_type: str,
    items: Iterable[dict],
    limit: int = 10,
    board_url: str = "",
    progress_callback: ProgressCallback | None = None,
    progress_state: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    saved = 0
    updated = 0
    failed: List[str] = []
    board_items = list(items)[:limit]
    for index, item in enumerate(board_items, start=1):
        bvid = str(item.get("bvid") or "").strip()
        if not bvid:
            continue
        detail = _video_detail(bvid) or dict(item)
        title = str(detail.get("title") or item.get("title") or "").strip()
        aid = safe_int(detail.get("aid") or item.get("aid") or 0)
        tags = _video_tags(detail or item)
        hotwords = _comment_hotwords(aid) if aid > 0 else []
        text = _structured_video_text(board_type, item, detail, tags, hotwords, board_url=board_url)
        try:
            document_id = f"{board_type}:{bvid}"
            existed = document_exists(metadata_filter={"source": "bilibili_hot_sync", "board_type": board_type, "bvid": bvid})
            if existed:
                delete_documents(metadata_filter={"source": "bilibili_hot_sync", "board_type": board_type, "bvid": bvid})
            result = add_document(
                Document(
                    id=document_id,
                    text=text,
                    metadata={
                        "source": "bilibili_hot_sync",
                        "board_type": board_type,
                        "bvid": bvid,
                        "title": title,
                        "board_url": str(board_url or "").strip(),
                        "partition": str(detail.get("tname") or item.get("tname") or "").strip(),
                    },
                )
            )
            saved += 1
            if existed or result.get("status") == "updated":
                updated += 1
        except Exception as exc:
            failed.append(f"{board_type}:{bvid}:{exc}")
        finally:
            if progress_state is not None:
                progress_state["processed_units"] = safe_int(progress_state.get("processed_units") or 0) + 1
                progress_state["processed_items"] = safe_int(progress_state.get("processed_items") or 0) + 1
                total_units = max(1, safe_int(progress_state.get("total_units") or 1))
                percent = min(99.4, (safe_int(progress_state.get("processed_units") or 0) / total_units) * 100)
                notify_progress(
                    progress_callback,
                    {
                        "status": "running",
                        "stage": "syncing_items",
                        "percent": round(percent, 2),
                        "message": f"正在同步 {board_type} 第 {index}/{len(board_items)} 条样本",
                        "board_type": board_type,
                        "current_title": title,
                        "board_item_index": index,
                        "board_item_total": len(board_items),
                        "processed_items": safe_int(progress_state.get("processed_items") or 0),
                        "total_items": safe_int(progress_state.get("total_items") or 0),
                        "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
                        "total_boards": safe_int(progress_state.get("total_boards") or 0),
                    },
                )
    return {"board_type": board_type, "saved_count": saved, "updated_count": updated, "failed": failed}


def crawl_and_store_bilibili_hot_videos(
    per_board_limit: int = 10,
    progress_callback: ProgressCallback | None = None,
) -> Dict[str, Any]:
    summary: List[Dict[str, Any]] = []
    board_total = 3 + len(POPULAR_RANK_PARTITIONS)
    board_item_limit = min(per_board_limit, 10)
    total_items = (per_board_limit * 3) + (board_item_limit * len(POPULAR_RANK_PARTITIONS))
    progress_state = {
        "processed_units": 0,
        "total_units": max(1, board_total + total_items),
        "processed_items": 0,
        "total_items": total_items,
        "processed_boards": 0,
        "total_boards": board_total,
    }

    notify_progress(
        progress_callback,
        {
            "status": "running",
            "stage": "prepare",
            "percent": 0.0,
            "message": "正在准备热门知识库更新任务",
            "processed_items": 0,
            "total_items": total_items,
            "processed_boards": 0,
            "total_boards": board_total,
        },
    )

    for legacy_partition in LEGACY_PRIMARY_PARTITIONS:
        try:
            delete_documents(metadata_filter={"source": "bilibili_hot_sync", "board_type": f"分区热门榜:{legacy_partition}"})
        except Exception:
            pass

    notify_progress(
        progress_callback,
        {
            "status": "running",
            "stage": "fetching_boards",
            "percent": 1.0,
            "message": "正在抓取全站热门榜",
            "board_type": "全站热门榜",
            "processed_items": 0,
            "total_items": total_items,
            "processed_boards": 0,
            "total_boards": board_total,
        },
    )
    hot_payload = _safe_sync(hot.get_hot_videos(ps=per_board_limit, pn=1), {})
    hot_items = hot_payload if isinstance(hot_payload, list) else hot_payload.get("list", []) or []
    progress_state["processed_units"] = safe_int(progress_state.get("processed_units") or 0) + 1
    summary.append(
        _ingest_hot_items(
            "全站热门榜",
            hot_items,
            limit=per_board_limit,
            board_url="https://www.bilibili.com/v/popular/rank/all",
            progress_callback=progress_callback,
            progress_state=progress_state,
        )
    )
    progress_state["processed_boards"] = safe_int(progress_state.get("processed_boards") or 0) + 1
    notify_progress(
        progress_callback,
        {
            "status": "running",
            "stage": "board_complete",
            "percent": round(min(99.4, (safe_int(progress_state.get("processed_units") or 0) / max(1, safe_int(progress_state.get("total_units") or 1))) * 100), 2),
            "message": "全站热门榜同步完成",
            "board_type": "全站热门榜",
            "processed_items": safe_int(progress_state.get("processed_items") or 0),
            "total_items": total_items,
            "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
            "total_boards": board_total,
        },
    )

    weekly_series = _safe_sync(hot.get_weekly_hot_videos_list(), {})
    weekly_list = (weekly_series.get("list") or weekly_series.get("data") or []) if isinstance(weekly_series, dict) else []
    latest_week = int((weekly_list[0] or {}).get("number") or 1) if weekly_list else 1
    notify_progress(
        progress_callback,
        {
            "status": "running",
            "stage": "fetching_boards",
            "percent": round(min(99.4, (safe_int(progress_state.get("processed_units") or 0) / max(1, safe_int(progress_state.get("total_units") or 1))) * 100), 2),
            "message": "正在抓取每周必看",
            "board_type": "每周必看",
            "processed_items": safe_int(progress_state.get("processed_items") or 0),
            "total_items": total_items,
            "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
            "total_boards": board_total,
        },
    )
    weekly_payload = _safe_sync(hot.get_weekly_hot_videos(latest_week), {})
    weekly_items = weekly_payload.get("list", []) if isinstance(weekly_payload, dict) else []
    progress_state["processed_units"] = safe_int(progress_state.get("processed_units") or 0) + 1
    summary.append(
        _ingest_hot_items(
            "每周必看",
            weekly_items,
            limit=per_board_limit,
            board_url="https://www.bilibili.com/v/popular/weekly/",
            progress_callback=progress_callback,
            progress_state=progress_state,
        )
    )
    progress_state["processed_boards"] = safe_int(progress_state.get("processed_boards") or 0) + 1
    notify_progress(
        progress_callback,
        {
            "status": "running",
            "stage": "board_complete",
            "percent": round(min(99.4, (safe_int(progress_state.get("processed_units") or 0) / max(1, safe_int(progress_state.get("total_units") or 1))) * 100), 2),
            "message": "每周必看同步完成",
            "board_type": "每周必看",
            "processed_items": safe_int(progress_state.get("processed_items") or 0),
            "total_items": total_items,
            "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
            "total_boards": board_total,
        },
    )

    notify_progress(
        progress_callback,
        {
            "status": "running",
            "stage": "fetching_boards",
            "percent": round(min(99.4, (safe_int(progress_state.get("processed_units") or 0) / max(1, safe_int(progress_state.get("total_units") or 1))) * 100), 2),
            "message": "正在抓取入站必刷",
            "board_type": "入站必刷",
            "processed_items": safe_int(progress_state.get("processed_items") or 0),
            "total_items": total_items,
            "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
            "total_boards": board_total,
        },
    )
    history_payload = _safe_sync(hot.get_history_popular_videos(), {})
    history_items = history_payload.get("list", []) if isinstance(history_payload, dict) else []
    progress_state["processed_units"] = safe_int(progress_state.get("processed_units") or 0) + 1
    summary.append(
        _ingest_hot_items(
            "入站必刷",
            history_items,
            limit=per_board_limit,
            board_url="https://www.bilibili.com/v/popular/history",
            progress_callback=progress_callback,
            progress_state=progress_state,
        )
    )
    progress_state["processed_boards"] = safe_int(progress_state.get("processed_boards") or 0) + 1
    notify_progress(
        progress_callback,
        {
            "status": "running",
            "stage": "board_complete",
            "percent": round(min(99.4, (safe_int(progress_state.get("processed_units") or 0) / max(1, safe_int(progress_state.get("total_units") or 1))) * 100), 2),
            "message": "入站必刷同步完成",
            "board_type": "入站必刷",
            "processed_items": safe_int(progress_state.get("processed_items") or 0),
            "total_items": total_items,
            "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
            "total_boards": board_total,
        },
    )

    for partition in POPULAR_RANK_PARTITIONS:
        board_type = f"分区热门榜:{partition.get('label')}"
        notify_progress(
            progress_callback,
            {
                "status": "running",
                "stage": "fetching_boards",
                "percent": round(min(99.4, (safe_int(progress_state.get("processed_units") or 0) / max(1, safe_int(progress_state.get("total_units") or 1))) * 100), 2),
                "message": f"正在抓取 {board_type}",
                "board_type": board_type,
                "processed_items": safe_int(progress_state.get("processed_items") or 0),
                "total_items": total_items,
                "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
                "total_boards": board_total,
            },
        )
        top_payload = _safe_sync(video_zone.get_zone_top10(int(partition.get("tid") or 0)), {})
        top_items = top_payload if isinstance(top_payload, list) else top_payload.get("list", []) or top_payload.get("top", []) or []
        progress_state["processed_units"] = safe_int(progress_state.get("processed_units") or 0) + 1
        summary.append(
            _ingest_hot_items(
                board_type,
                top_items,
                limit=min(per_board_limit, 10),
                board_url=str(partition.get("board_url") or "https://www.bilibili.com/v/popular/rank/"),
                progress_callback=progress_callback,
                progress_state=progress_state,
            )
        )
        progress_state["processed_boards"] = safe_int(progress_state.get("processed_boards") or 0) + 1
        notify_progress(
            progress_callback,
            {
                "status": "running",
                "stage": "board_complete",
                "percent": round(min(99.4, (safe_int(progress_state.get("processed_units") or 0) / max(1, safe_int(progress_state.get("total_units") or 1))) * 100), 2),
                "message": f"{board_type} 同步完成",
                "board_type": board_type,
                "processed_items": safe_int(progress_state.get("processed_items") or 0),
                "total_items": total_items,
                "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
                "total_boards": board_total,
            },
        )

    result = {
        "status": "ok",
        "boards": summary,
        "total_saved": sum(item.get("saved_count", 0) for item in summary),
        "total_updated": sum(item.get("updated_count", 0) for item in summary),
        "total_failed": sum(len(item.get("failed", [])) for item in summary),
    }
    notify_progress(
        progress_callback,
        {
            "status": "completed",
            "stage": "completed",
            "percent": 100.0,
            "message": "热门知识库更新完成",
            "processed_items": safe_int(progress_state.get("processed_items") or 0),
            "total_items": total_items,
            "processed_boards": safe_int(progress_state.get("processed_boards") or 0),
            "total_boards": board_total,
            "result": result,
        },
    )
    return result


def update_chroma_knowledge_base(
    per_board_limit: int = 10,
    progress_callback: ProgressCallback | None = None,
) -> Dict[str, Any]:
    if progress_callback is None:
        return crawl_and_store_bilibili_hot_videos(per_board_limit=per_board_limit)
    return crawl_and_store_bilibili_hot_videos(per_board_limit=per_board_limit, progress_callback=progress_callback)
